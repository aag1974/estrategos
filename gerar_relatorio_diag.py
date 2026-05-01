"""
gerar_relatorio_diag.py — Gerador Word do Diagnóstico (tipo Reposicionamento).

Uso:
  python3 gerar_relatorio_diag.py diag_config.json

Entrada (JSON):
  {
    "tipo": "reposicionamento",
    "origem":     {"nome": "MANZONI MARQUES", "cargo": "DEPUTADO_DISTRITAL"},
    "destino":    "DEPUTADO_FEDERAL",
    "referencia": {"nome": "BIA KICIS",      "cargo": "DEPUTADO_FEDERAL"}
  }

Saída:
  relatorio_diag_<slug>.docx — autocontido, com imagens embarcadas.
"""

import json, sys, re, tempfile, shutil
from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.colors import LinearSegmentedColormap, Normalize
import geopandas as gpd

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from fase4_v2 import carregar, montar_dados, montar_candidatos

# ── Constantes ────────────────────────────────────────────────────────────────

GEOJSON_PATH = Path("Limite_RA_20190.json")

CARGO_LBL = {
    "GOVERNADOR": "Governador",
    "SENADOR": "Senador",
    "DEPUTADO_FEDERAL": "Deputado Federal",
    "DEPUTADO_DISTRITAL": "Deputado Distrital",
}
CAMPO_LBL = {
    "progressista": "Progressista",
    "moderado": "Moderado",
    "liberal_conservador": "Liberal/Conservador",
    "outros": "Outros",
}
META_CARGO = {
    "GOVERNADOR": 700_000,
    "SENADOR": 550_000,
    "DEPUTADO_FEDERAL": 18_000,
    "DEPUTADO_DISTRITAL": 18_000,
}
SEM_ZONA = {"Park Way", "SIA", "Fercal", "Sol Nascente/Pôr do Sol", "Arniqueira"}

CAT_DEF = {
    "compartilhado": {"lbl": "Base compartilhada", "cor": "#0F6E56",
        "desc": "Regiões onde o candidato e a referência foram fortes (Performance positiva nos dois). Voto naturalmente compartilhado — eleitor compatível com os dois perfis."},
    "nucleo": {"lbl": "Voto pessoal", "cor": "#3D7BAB",
        "desc": "Regiões onde só o candidato foi forte (a referência foi fraca). Voto puxado pelo nome do candidato, não pela base do cargo destino."},
    "lacuna": {"lbl": "Espaço a conquistar", "cor": "#B45309",
        "desc": "Regiões onde só a referência foi forte (o candidato não chegou). Eleitorado da referência ainda inacessado pelo candidato."},
    "hostil": {"lbl": "Terreno aberto", "cor": "#A32D2D",
        "desc": "Regiões onde nem o candidato nem a referência foram fortes, mas com volume eleitoral alto. Eleitor numeroso, sem dono claro."},
    "periferia": {"lbl": "Volume baixo", "cor": "#6B7280",
        "desc": "Regiões onde nem o candidato nem a referência foram fortes, e o volume eleitoral é pequeno. Pouca gente, pouco voto."},
}
CAT_ORDER = ["compartilhado", "nucleo", "lacuna", "hostil", "periferia"]


# ── Formatadores pt-BR ────────────────────────────────────────────────────────

def fmt_int(n):
    if n is None: return "—"
    return f"{int(round(n)):,}".replace(",", ".")

def fmt_pct(v, dec=1):
    if v is None: return "—"
    s = f"{v:.{dec}f}".replace(".", ",")
    return s + "%"

def fmt_idx(v):
    if v is None: return "—"
    d = round((v - 1) * 100)
    s = ("+" if d >= 0 else "") + f"{d}%"
    return s


# ── Tradução socioterritorial ─────────────────────────────────────────────────
# Classifica cada RA por perfil socioeconômico para que a análise possa dizer
# coisas como "base periférica popular" em vez de listar nomes de RAs solto.

PERFIL_RA = {
    # Cinturão AB consolidado
    "Plano Piloto": "AB", "Brasília (Plano Piloto)": "AB",
    "Lago Sul": "AB", "Lago Norte": "AB",
    "Sudoeste/Octogonal": "AB", "Jardim Botânico": "AB", "Park Way": "AB",
    # Classe média urbana
    "Águas Claras": "media", "Taguatinga": "media", "Guará": "media",
    "Sobradinho": "media", "Sobradinho II": "media", "Cruzeiro": "media",
    "Núcleo Bandeirante": "media", "Vicente Pires": "media",
    "Candangolândia": "media", "Riacho Fundo": "media",
    "Arniqueira": "media", "SIA": "media",
    # Periferia popular
    "Ceilândia": "popular", "Samambaia": "popular", "Santa Maria": "popular",
    "Recanto das Emas": "popular", "Riacho Fundo II": "popular",
    "Paranoá": "popular", "Itapoã": "popular", "SCIA/Estrutural": "popular",
    "Sol Nascente/Pôr do Sol": "popular", "Varjão": "popular", "Estrutural": "popular",
    # Interior agro / periferia distante
    "Brazlândia": "rural", "Planaltina": "rural",
    "São Sebastião": "rural", "Gama": "rural", "Fercal": "rural",
}

PERFIL_DESC = {
    "AB":      "regiões de maior renda do DF",
    "media":   "satélites de classe média",
    "popular": "periferias populares",
    "rural":   "regiões do interior do DF",
}
PERFIL_BASE = {
    "AB":      "fortalezas nas regiões de maior renda",
    "media":   "fortalezas em satélites de classe média",
    "popular": "fortalezas nas periferias populares",
    "rural":   "fortalezas em regiões do interior",
}

def perfil_dominante(ras_list):
    """Retorna (perfil_dominante, pct_dominancia) entre AB/media/popular/rural."""
    contagens = {}
    for ra in ras_list:
        p = PERFIL_RA.get(ra, "outro")
        contagens[p] = contagens.get(p, 0) + 1
    if not contagens: return (None, 0)
    dom = max(contagens, key=contagens.get)
    return (dom, contagens[dom] / sum(contagens.values()))

def descrever_base(ras_strong):
    """Frase tipo 'base periférica popular' com base nos top redutos."""
    p, pct = perfil_dominante(ras_strong)
    if not p: return "base sem perfil definido"
    if pct >= 0.66: return PERFIL_BASE.get(p, "base mista")
    return "base de perfil misto"

def descrever_ausencias(ras_weak):
    """Frase tipo 'cinturão classe AB' a partir das maiores ausências."""
    p, pct = perfil_dominante(ras_weak)
    if not p: return "regiões dispersas"
    if pct >= 0.66: return PERFIL_DESC.get(p, "regiões mistas")
    return "regiões de perfil misto"


# ── Carregamento ──────────────────────────────────────────────────────────────

def slugify(s):
    s = re.sub(r"[^\w\s-]", "", s, flags=re.UNICODE).strip().lower()
    s = re.sub(r"[\s_-]+", "_", s)
    return s

def buscar_candidato(cands, nome, cargo):
    nome_u = nome.upper().strip()
    for c in cands:
        if c["cargo"] == cargo and c["nome"].upper().strip() == nome_u:
            return c
    # Fallback fuzzy: contém
    for c in cands:
        if c["cargo"] == cargo and nome_u in c["nome"].upper():
            return c
    raise ValueError(f"Candidato não encontrado: {nome} / {cargo}")


# ── Cálculo das 5 categorias ──────────────────────────────────────────────────

def computar_categorias(origem, ref, dados_ra):
    ras = [n for n in dados_ra.keys() if n not in SEM_ZONA]
    aptos_arr = sorted([(dados_ra[n].get("el_aptos") or 0) for n in ras])
    med_aptos = aptos_arr[len(aptos_arr) // 2] if aptos_arr else 0

    cats = {k: [] for k in CAT_ORDER}
    for n in ras:
        ro = (origem.get("ras") or {}).get(n)
        rr = (ref.get("ras") or {}).get(n)
        ido = ro.get("idx") if ro else None
        idr = rr.get("idx") if rr else None
        f_orig = ido is not None and ido >= 1.0
        f_ref = idr is not None and idr >= 1.0
        aptos = dados_ra[n].get("el_aptos") or 0
        item = {
            "ra": n, "aptos": aptos,
            "orig_votos": (ro.get("v") if ro else 0) or 0,
            "ref_votos":  (rr.get("v") if rr else 0) or 0,
            "orig_idx": ido, "ref_idx": idr,
        }
        if f_orig and f_ref: cats["compartilhado"].append(item)
        elif f_orig:         cats["nucleo"].append(item)
        elif f_ref:          cats["lacuna"].append(item)
        elif aptos >= med_aptos: cats["hostil"].append(item)
        else:                cats["periferia"].append(item)

    for k in cats:
        cats[k].sort(key=lambda r: -r["aptos"])
    return cats, med_aptos, len(ras)


def calcular_tipologia(cand):
    deltas = []
    for ra, v in (cand.get("ras") or {}).items():
        idx = v.get("idx")
        if idx is not None:
            deltas.append((idx - 1) * 100)
    if len(deltas) < 2: return None
    media = sum(deltas) / len(deltas)
    sigma = (sum((d - media) ** 2 for d in deltas) / len(deltas)) ** 0.5
    if sigma < 30:   lbl = "Distribuído"
    elif sigma < 60: lbl = "Híbrido"
    else:            lbl = "Concentrado"
    return {"lbl": lbl, "sigma": sigma}


# ── Geração de imagens ────────────────────────────────────────────────────────

def gerar_scatter_png(dados_ra, out_path):
    pontos = []
    for nome, d in dados_ra.items():
        sf = d.get("pct_serv_fed")
        ab = d.get("pct_ab")
        v = d.get("votos", {}).get("DEPUTADO_FEDERAL", {}).get("progressista", {})
        pct = v.get("pct") if isinstance(v, dict) else None
        if sf is None or pct is None or ab is None: continue
        pontos.append({"ra": nome, "x": float(sf), "y": float(pct), "ab": float(ab)})
    if len(pontos) < 5: return False

    xs = np.array([p["x"] for p in pontos])
    ys = np.array([p["y"] for p in pontos])
    abs_ = np.array([p["ab"] for p in pontos])
    a, b = np.polyfit(xs, ys, 1)
    r = np.corrcoef(xs, ys)[0, 1]

    fig, ax = plt.subplots(figsize=(7, 4), dpi=150)
    cmap = LinearSegmentedColormap.from_list("ab", ["#D2D2D2", "#534AB7"])
    sc = ax.scatter(xs, ys, c=abs_, cmap=cmap, s=70, edgecolors="white", linewidths=0.8, zorder=3)
    xline = np.array([xs.min(), xs.max()])
    ax.plot(xline, a * xline + b, "--", color="#534AB7", alpha=0.6, linewidth=1.4, zorder=2)
    ax.set_xlabel("% Servidor federal na população (PDAD)", fontsize=10)
    ax.set_ylabel("% Voto progressista (Federal 2022)", fontsize=10)
    # Sem título embarcado — entra como caption no Word
    ax.text(0.97, 0.95, f"r = {r:+.2f}".replace(".", ","),
            transform=ax.transAxes, ha="right", va="top",
            fontsize=12, color="#534AB7", fontweight="bold")
    ax.grid(True, color="#f1f1f1", linewidth=0.5)
    ax.set_axisbelow(True)
    cbar = plt.colorbar(sc, ax=ax, fraction=0.04, pad=0.02)
    cbar.set_label("% Classe AB", fontsize=9)
    cbar.ax.tick_params(labelsize=8)
    # Sem título embarcado — entra como legenda no Word
    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return True


def carregar_geo():
    gdf = gpd.read_file(GEOJSON_PATH)
    if "ra" not in gdf.columns:
        for c in gdf.columns:
            if c.lower() == "ra":
                gdf = gdf.rename(columns={c: "ra"}); break
    return gdf


def gerar_mapa_idx(cand, gdf, out_path, titulo):
    """Mapa das RAs coloridas por sobre-índice (idx-1)*100."""
    deltas = {}
    for n, v in (cand.get("ras") or {}).items():
        idx = v.get("idx")
        if idx is not None:
            deltas[n] = (idx - 1) * 100

    g = gdf.copy()
    g["delta"] = g["ra"].map(deltas)

    # Paleta divergente (vermelho → cinza → verde) centrada em 0
    cmap = LinearSegmentedColormap.from_list(
        "div", ["#A32D2D", "#FCEBEB", "#F1EFE8", "#E1F5EE", "#0F6E56"])
    vmax = max(60, abs(g["delta"].dropna().abs().max() or 0))
    norm = Normalize(vmin=-vmax, vmax=vmax)

    fig, ax = plt.subplots(figsize=(7.5, 6), dpi=150)
    g.plot(column="delta", cmap=cmap, norm=norm, ax=ax,
           edgecolor="#fff", linewidth=0.6,
           missing_kwds={"color": "#E5E5E5", "edgecolor": "#fff", "linewidth": 0.4})
    # Sem título embarcado — entra como caption no Word
    ax.axis("off")
    sm = plt.cm.ScalarMappable(cmap=cmap, norm=norm); sm.set_array([])
    cbar = plt.colorbar(sm, ax=ax, fraction=0.035, pad=0.02, orientation="vertical")
    cbar.set_label("Performance (%)", fontsize=9)
    cbar.ax.tick_params(labelsize=8)
    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def gerar_mapa_categorias(cats, gdf, out_path):
    """Mapa das RAs coloridas pelas 5 categorias do Reposicionamento."""
    from matplotlib.patches import Patch

    cat_por_ra = {}
    for k in CAT_ORDER:
        for r in cats[k]:
            cat_por_ra[r["ra"]] = k

    g = gdf.copy()
    g["cat"] = g["ra"].map(cat_por_ra)

    color_map = {k: CAT_DEF[k]["cor"] for k in CAT_ORDER}

    fig, ax = plt.subplots(figsize=(8, 6), dpi=150)
    legend_handles = []
    for k in CAT_ORDER:
        sub = g[g["cat"] == k]
        if not sub.empty:
            sub.plot(ax=ax, color=color_map[k], edgecolor="#fff", linewidth=0.6)
            legend_handles.append(Patch(facecolor=color_map[k], edgecolor="#fff",
                                        label=f"{CAT_DEF[k]['lbl']} ({len(cats[k])})"))
    miss = g[g["cat"].isna()]
    if not miss.empty:
        miss.plot(ax=ax, color="#E5E5E5", edgecolor="#fff", linewidth=0.4)
        legend_handles.append(Patch(facecolor="#E5E5E5", edgecolor="#fff", label="Sem dado"))
    # Sem título embarcado — entra como caption no Word
    ax.axis("off")
    ax.legend(handles=legend_handles, loc="lower left", fontsize=8, frameon=False)
    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


# ── Helpers python-docx ───────────────────────────────────────────────────────
# Estilos derivados do template Etapa1_Cliente_Secreto.docx:
#   Heading 1: 16pt bold, cor #004C99 (azul corporativo)
#   Heading 2: 13pt bold, cor #747678 (cinza médio)
#   Heading 3: bold, ~12pt
#   Body:      Calibri 11pt
#   Captions:  bold pequeno (Tabela acima, Figura abaixo)

COR_AZUL = RGBColor(0x00, 0x4C, 0x99)
COR_CINZA = RGBColor(0x74, 0x76, 0x78)
COR_TXT = RGBColor(0x1A, 0x1A, 0x1A)
COR_MUTED = RGBColor(0x6B, 0x72, 0x80)


def setup_styles(doc):
    """Aplica os estilos do template Etapa1 nos estilos nativos do documento."""
    n = doc.styles["Normal"]
    n.font.name = "Calibri"; n.font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16); h1.font.bold = True; h1.font.color.rgb = COR_AZUL

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13); h2.font.bold = True; h2.font.color.rgb = COR_CINZA

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12); h3.font.bold = True; h3.font.color.rgb = COR_TXT


def add_h1(doc, txt):
    p = doc.add_paragraph(txt, style="Heading 1")
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_h2(doc, txt):
    p = doc.add_paragraph(txt, style="Heading 2")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_h3(doc, txt):
    p = doc.add_paragraph(txt, style="Heading 3")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_kicker(doc, txt):
    """Linha de etiqueta no topo da capa (uppercase, pequeno)."""
    p = doc.add_paragraph()
    r = p.add_run(txt.upper())
    r.font.size = Pt(9); r.font.color.rgb = COR_MUTED; r.bold = True
    r.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)

def add_p(doc, txt):
    """Parágrafo de corpo com suporte a **negrito**."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    parts = re.split(r"(\*\*[^*]+\*\*)", txt)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        else:
            p.add_run(part)
    return p

def add_block_note(doc, label, txt):
    """Caixa de nota metodológica (Block Text) — destaque visual leve."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    if label:
        r1 = p.add_run(label + " ")
        r1.bold = True; r1.font.color.rgb = COR_AZUL; r1.font.size = Pt(10.5)
    parts = re.split(r"(\*\*[^*]+\*\*)", txt)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            rr = p.add_run(part[2:-2]); rr.bold = True; rr.font.size = Pt(10.5)
        else:
            rr = p.add_run(part); rr.font.size = Pt(10.5)
    return p

# Contadores de figura/tabela (usar via FIG_N[0] / TAB_N[0] em montar_word)
def reset_counters():
    return {"fig": [0], "tab": [0]}

def add_caption_tabela(doc, ctr, titulo):
    """Caption acima da tabela: 'Tabela N – Título' bold."""
    ctr["tab"][0] += 1
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Tabela {ctr['tab'][0]} – {titulo}")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = COR_TXT
    return p

def add_caption_figura(doc, ctr, titulo):
    """Caption abaixo da figura: 'Figura N – Título' bold, centralizado."""
    ctr["fig"][0] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(f"Figura {ctr['fig'][0]} – {titulo}")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = COR_TXT
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True; run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            para = cells[i].paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in t.columns[i].cells:
                cell.width = w
    return t


# ── Helpers de leitura narrativa (parametrizam o texto) ───────────────────────

def implicacao_paradoxo_campo(campo, cargo_dest):
    """Lê o paradoxo (regiões ricas votam progressista) sob a ótica do campo do candidato. ELI5."""
    if campo == "progressista":
        base = ("Para um candidato do campo progressista no DF, isso é uma vantagem natural. "
                "As regiões de maior renda — Plano Piloto, Lago Sul, Sudoeste, Jardim Botânico, "
                "Águas Claras — costumam votar nesse campo. O ponto de atenção é que essas "
                "regiões, juntas, têm menos eleitores do que as periferias do DF, então a base "
                "natural tem teto em volume. ")
    elif campo == "moderado":
        base = ("Para um candidato do campo moderado, o terreno é misto. O moderado dialoga "
                "tanto com a classe média urbana quanto com as periferias e satélites consolidadas "
                "(Taguatinga, Guará, Sobradinho), mas compete com progressistas nas regiões mais "
                "ricas e com liberais-conservadores nas periferias populares. ")
    elif campo == "liberal_conservador":
        base = ("Para um candidato do campo liberal-conservador, as regiões mais ricas do DF "
                "tendem a ser terreno difícil — não é só com este candidato, é uma característica "
                "do DF. As bases naturais desse campo costumam estar nas periferias populares "
                "(Ceilândia, Santa Maria, Samambaia) e na classe média periférica (Taguatinga, "
                "Guará). ")
    else:
        base = ""
    if cargo_dest == "GOVERNADOR":
        base += ("Em particular, na disputa para Governador, o peso da máquina política local "
                 "e da rede de servidores distritais costuma ser determinante — mais do que o "
                 "alinhamento ideológico puro. É o cargo mais sensível à gestão e a vínculos "
                 "concretos com cada região.")
    elif cargo_dest in ("DEPUTADO_FEDERAL", "SENADOR"):
        base += ("Em particular, no cargo de " + CARGO_LBL.get(cargo_dest, cargo_dest) +
                 ", a influência da classe média alta e do servidor federal aparece com mais "
                 "força na composição do voto.")
    elif cargo_dest == "DEPUTADO_DISTRITAL":
        base += ("Em particular, no Distrital, o que mais pesa são vínculos territoriais "
                 "concretos com cada região — máquina local, rede de bairro, presença diária. "
                 "O efeito do perfil socioeconômico é mais fraco do que nos cargos federais.")
    return base


def descrever_perfil_votacao(tipologia, nome):
    """
    Traduz Distribuído/Híbrido/Concentrado em frase prática,
    sem termo estatístico (σ pp).
    """
    if not tipologia: return ""
    lbl = tipologia["lbl"]
    if lbl == "Distribuído":
        return (f"O **Perfil de votação** de {nome} é **Distribuído**. A votação dele se "
                f"espalha de forma parecida pelas regiões — sem fortalezas marcantes, mas "
                f"com presença em quase todo lugar. É um candidato de capilaridade ampla, com "
                f"base larga em vez de profunda.")
    if lbl == "Híbrido":
        return (f"O **Perfil de votação** de {nome} é **Híbrido**. Tem fortalezas claras em "
                f"algumas regiões, mas chega em todas as outras com presença razoável. É um "
                f"candidato com reduto definido que ainda assim alcança o resto do território.")
    if lbl == "Concentrado":
        return (f"O **Perfil de votação** de {nome} é **Concentrado**. O voto se prende a "
                f"poucas regiões — em algumas tem presença forte, em outras quase desaparece. "
                f"É um candidato de nicho territorial: fora do reduto, a votação cai bastante.")
    return ""

# Compatibilidade com chamadas antigas
def descrever_tipologia(tipologia, nome):
    return descrever_perfil_votacao(tipologia, nome)


def leitura_cenarios(pct_cns, lacuna_pct_meta):
    """
    Leitura neutra (sem veredicto) sobre o que o cenário máximo indica.
    Descreve o que os números mostram, sem julgar viabilidade.
    """
    if pct_cns >= 90:
        return ("No cenário máximo do modelo (Ponte construída), a captura integral do Espaço "
                "a conquistar pela referência cobre a totalidade ou quase a totalidade do "
                "Patamar de eleição. Os números indicam que o eleitorado da referência, se "
                "capturado em proporção elevada, é suficiente para colocar a candidatura na "
                "faixa de eleição — sem necessidade de expansão fora dela.")
    if pct_cns >= 60:
        return ("No cenário máximo do modelo, a captura do Espaço a conquistar cobre cerca de "
                "dois terços ou mais do Patamar de eleição. Os números indicam que o eleitorado "
                "da referência cobre a maior parte da diferença até a meta, mas não toda — "
                "uma parcela do volume necessário fica fora do escopo medido por este modelo.")
    if pct_cns >= 35:
        return ("No cenário máximo do modelo, a captura do Espaço a conquistar cobre menos da "
                "metade do Patamar de eleição. Os números indicam que parte significativa do "
                "volume necessário está fora do escopo medido por este modelo — vem de outras "
                "fontes (Terreno aberto, alianças, ou outro cargo destino).")
    return ("No cenário máximo do modelo, a captura integral do Espaço a conquistar cobre uma "
            "parcela pequena do Patamar de eleição. Os números indicam que a maior parte do "
            "volume necessário está fora do escopo medido por este modelo — vem de territórios "
            "onde nem o candidato nem a referência foram fortes, de alianças com perfis "
            "complementares, ou da escolha de outro cargo destino. Como pesar essas opções "
            "é decisão da campanha.")


def perfil_intersecao(cats):
    """Frase qualitativa sobre tamanho da interseção candidato × referência (ELI5)."""
    n_comp = len(cats["compartilhado"])
    n_lac = len(cats["lacuna"])
    if n_comp >= 10 and n_lac <= 4:
        return ("Candidato e referência têm muitas regiões em comum onde ambos foram fortes — "
                "estão na mesma vizinhança eleitoral. Por outro lado, sobra pouco eleitorado "
                "da referência fora dessa sobreposição. Ou seja: o eleitor é parecido, mas há "
                "pouco terreno novo onde apenas a referência foi forte.")
    if n_comp <= 4 and n_lac >= 8:
        return ("Candidato e referência têm pouca sobreposição direta — foram fortes em regiões "
                "diferentes. Sobra muito eleitorado da referência ainda não acessado pelo "
                "candidato. Ou seja: há terreno novo amplo, mas a campanha começa partindo de "
                "longe da referência.")
    if n_comp >= 8 and n_lac >= 6:
        return ("Candidato e referência têm boa sobreposição em várias regiões e ainda sobra "
                "eleitorado da referência fora dessa interseção. É uma combinação confortável: "
                "base segura de partida e espaço relevante para crescer pela referência.")
    return ("Candidato e referência têm sobreposição em parte das regiões e diferenças em "
            "outra parte — escala intermediária, sem dominância clara nem do compartilhado "
            "nem do espaço a conquistar.")


# ── Construção do Word ────────────────────────────────────────────────────────



# ── Construção do Word ────────────────────────────────────────────────────────

def montar_word(config, origem, ref, dados_ra, cats, med_aptos, total_ras,
                imgs, output_path):
    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)

    ctr = reset_counters()

    # ── Variáveis ────────────────────────────────────────────────────────────
    cargo_dest = config["destino"]
    cargo_dest_lbl = CARGO_LBL.get(cargo_dest, cargo_dest)
    orig_cargo_lbl = CARGO_LBL.get(origem["cargo"], origem["cargo"])
    nome = origem["nome"]
    nome_curto = nome.split()[0].title() if nome else nome
    ref_nome = ref["nome"]
    ref_nome_curto = ref_nome.split()[0].title() if ref_nome else ref_nome
    campo = origem.get("campo", "outros")
    campo_lbl = CAMPO_LBL.get(campo, campo)

    meta = META_CARGO.get(cargo_dest, 18000)
    substituto = sum(r["orig_votos"] for r in cats["compartilhado"]) + sum(r["orig_votos"] for r in cats["nucleo"])
    lacuna_total = sum(r["ref_votos"] for r in cats["lacuna"])
    pcl = substituto + 0.30 * lacuna_total
    cns = substituto + 0.60 * lacuna_total
    pct_sub = substituto / meta * 100 if meta > 0 else 0
    pct_pcl = pcl / meta * 100 if meta > 0 else 0
    pct_cns = cns / meta * 100 if meta > 0 else 0
    falta_meta = max(0, int(meta - cns))
    salto = meta / origem["total"] if origem["total"] > 0 else 0
    pct_origem_meta = origem["total"] / meta * 100 if meta else 0
    salto_str = f"{salto:.1f}".replace(".", ",")

    ras_origem = [{"ra": n, **v} for n, v in (origem.get("ras") or {}).items() if (v.get("v") or 0) > 0]
    forcas = sorted(ras_origem, key=lambda r: -(r.get("idx") or 0))[:3]
    fraquezas = sorted(ras_origem, key=lambda r: (r.get("idx") if r.get("idx") is not None else 99))[:3]

    ras_ref = [{"ra": n, **v} for n, v in (ref.get("ras") or {}).items() if (v.get("v") or 0) > 0]
    ref_forcas_idx = sorted(ras_ref, key=lambda r: -(r.get("idx") or 0))[:3]
    ref_forcas_vol = sorted(ras_ref, key=lambda r: -(r.get("v") or 0))[:3]
    ref_fraquezas = sorted(ras_ref, key=lambda r: (r.get("idx") if r.get("idx") is not None else 99))[:3]
    n_ras_ref_pos = sum(1 for r in ras_ref if r.get("idx") is not None and r["idx"] >= 1.0)

    tip_o = calcular_tipologia(origem)
    tip_r = calcular_tipologia(ref)

    arr_comp = cats["compartilhado"]
    arr_nuc = cats["nucleo"]
    arr_lac = cats["lacuna"]
    arr_hos = cats["hostil"]
    arr_per = cats["periferia"]
    votos_comp = sum(r["orig_votos"] for r in arr_comp)
    votos_nuc = sum(r["orig_votos"] for r in arr_nuc)
    aptos_hos = sum(r["aptos"] for r in arr_hos)
    aptos_per = sum(r["aptos"] for r in arr_per)

    perfil_forcas = perfil_dominante([r["ra"] for r in forcas])[0]
    perfil_fraq = perfil_dominante([r["ra"] for r in fraquezas])[0]
    perfil_lac = perfil_dominante([r["ra"] for r in arr_lac])[0]
    perfil_lac_desc = PERFIL_DESC.get(perfil_lac, "regiões mistas") if arr_lac else None
    perfil_ref_forcas = perfil_dominante([r["ra"] for r in ref_forcas_idx])[0]

    razao_ref_orig = ref["total"] / origem["total"] if origem["total"] > 0 else 0
    razao_ref_str = f"{razao_ref_orig:.1f}".replace(".", ",")

    # Comparação ref vs Patamar
    ref_pct_meta = ref["total"] / meta * 100 if meta else 0
    if ref["total"] >= meta * 1.10:
        ref_vs_meta = f"acima do Patamar de eleição (~{int(ref_pct_meta-100)}% além)"
    elif ref["total"] >= meta * 0.95:
        ref_vs_meta = "na faixa do Patamar de eleição"
    elif ref["total"] >= meta * 0.7:
        ref_vs_meta = f"próxima do Patamar de eleição ({fmt_pct(ref_pct_meta)} dele)"
    else:
        ref_vs_meta = f"abaixo do Patamar de eleição ({fmt_pct(ref_pct_meta)} dele)"

    # ── CAPA ──
    add_kicker(doc, "Diagnóstico Estratégico · Tipo Reposicionamento")
    add_h1(doc, "Relatório Técnico")

    cap_p = doc.add_paragraph()
    cap_p.paragraph_format.space_after = Pt(8)
    for line in [
        f"Tipo: Diagnóstico Estratégico — Reposicionamento Eleitoral",
        f"Candidato: {nome}",
        f"Cargo de origem (2022): {orig_cargo_lbl}",
        f"Cargo de destino (2026): {cargo_dest_lbl}",
        f"Referência destino: {ref_nome}",
        f"Período do dado: Eleições 2022 (TSE)",
    ]:
        r = cap_p.add_run(line + "\n")
        r.bold = True; r.font.size = Pt(11)

    add_h2(doc, "Ficha técnica")
    for line in [
        "Executora: Opinião Informação Estratégica",
        "Modelo: Diagnóstico Estratégico — Reposicionamento (DF)",
        "Bases analisadas: TSE 2022 (votos por seção, cadastro eleitoral); PDAD 2021 (perfil socioeconômico das Regiões Administrativas)",
        "Unidade de análise: Região Administrativa do DF (28 RAs com voto registrado)",
        "Software: Pipeline próprio (Python/pandas/geopandas)",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line); r.bold = True; r.font.size = Pt(10.5)

    doc.add_page_break()

    # ── 1. SUMÁRIO EXECUTIVO ──
    add_h1(doc, "1. Sumário Executivo")

    add_h2(doc, "1.1 Objetivo")
    add_p(doc,
        f"Avaliar como a votação de {nome} no cargo de {orig_cargo_lbl} em 2022 se "
        f"compara à votação de {ref_nome} no cargo de {cargo_dest_lbl} no mesmo ano, "
        f"com o objetivo de mapear o eleitorado já consolidado, o eleitorado "
        f"potencialmente capturável e a distância até o **Patamar de eleição** do cargo "
        f"que {nome_curto} pretende disputar em 2026.")

    add_h2(doc, "1.2 Escopo")
    add_p(doc,
        f"A análise considera as 28 Regiões Administrativas do DF com voto registrado no "
        f"TSE 2022. Para cada região, calculamos a **Performance** (peso do voto recebido "
        f"comparado à média do cargo) tanto do candidato quanto da referência. O cruzamento "
        f"dessas duas Performances classifica cada região em uma das **5 zonas territoriais** "
        f"do reposicionamento e alimenta os **3 cenários de captura**.")

    add_h2(doc, "1.3 Síntese")
    add_p(doc,
        f"{nome_curto} foi eleito {orig_cargo_lbl} em 2022 com **{fmt_int(origem['total'])} "
        f"votos**, alinhado ao campo {campo_lbl.lower()}. {ref_nome_curto} foi "
        f"{'eleito' if ref.get('eleito') else 'candidato'} {cargo_dest_lbl} em 2022 com "
        f"**{fmt_int(ref['total'])} votos**. A diferença entre os dois é de "
        f"**{razao_ref_str}x** — a referência tem essa razão de eleitorado em comparação "
        f"ao candidato, no DF.")
    add_p(doc,
        f"O **Patamar de eleição** para {cargo_dest_lbl} é **{fmt_int(meta)} votos**. "
        f"A votação atual de {nome_curto} cobre {fmt_pct(pct_origem_meta)} desse patamar.")

    add_caption_tabela(doc, ctr, "Resumo do cruzamento candidato × referência (cenários e zonas)")
    add_table(doc,
        ["Indicador", "Valor"],
        [
            ["Candidato", nome],
            ["Cargo de origem (2022)", f"{orig_cargo_lbl} · {fmt_int(origem['total'])} votos · {origem.get('partido','?')}"],
            ["Cargo de destino (2026)", cargo_dest_lbl],
            ["Referência destino", f"{ref_nome} · {fmt_int(ref['total'])} votos"],
            ["Patamar de eleição", fmt_int(meta) + " votos"],
            ["Cenário Substituto orgânico", f"{fmt_int(substituto)} ({fmt_pct(pct_sub)} do Patamar)"],
            ["Cenário Ponte parcial (30% do Espaço a conquistar)", f"{fmt_int(pcl)} ({fmt_pct(pct_pcl)} do Patamar)"],
            ["Cenário Ponte construída (60% do Espaço a conquistar)", f"{fmt_int(cns)} ({fmt_pct(pct_cns)} do Patamar)"],
            ["RAs em Base compartilhada", str(len(arr_comp))],
            ["RAs em Espaço a conquistar", str(len(arr_lac))],
        ])

    leitura = leitura_cenarios(pct_cns, lacuna_total / meta * 100 if meta else 0)
    add_block_note(doc, "O que os números mostram:", leitura)

    doc.add_page_break()

    # ── 2. NOTAS METODOLÓGICAS ──
    add_h1(doc, "2. Notas Metodológicas")
    add_p(doc,
        "Esta seção define os termos canônicos do produto. Quando aparecerem em **negrito** "
        "ao longo do relatório, estão definidos aqui.")

    add_h2(doc, "2.1 Performance")
    add_p(doc,
        "**Performance** é o quanto um candidato recebe de voto numa região comparado ao "
        "que seria esperado pelo tamanho do eleitorado dessa região. Performance **+50%** "
        "significa que recebeu 50% mais votos do que receberia se a votação fosse "
        "proporcional ao tamanho da RA. Performance **-50%** significa que recebeu metade "
        "do esperado. Performance **0%** significa que recebeu exatamente a média do cargo. "
        "É um indicador de força relativa territorial — não de volume absoluto.")

    add_h2(doc, "2.2 Patamar de eleição")
    add_p(doc,
        "**Patamar de eleição** é o volume de votos que candidatos individuais costumam "
        "alcançar para serem eleitos no cargo, com base nos resultados de 2022. Valores "
        "canônicos:")
    add_p(doc,
        "**Governador**: ~700.000 votos (mínimo histórico para chegar ao 2º turno). "
        "**Senador**: ~550.000 votos. **Deputado Federal e Distrital**: ~18.000 votos "
        "(faixa onde candidatos individuais começam a ser efetivamente eleitos). É uma "
        "referência empírica, não um cálculo formal.")

    add_h2(doc, "2.3 Perfil de votação")
    add_p(doc,
        "Descreve como a votação de um candidato se distribui entre as regiões. Três tipos:")
    add_p(doc,
        "**Distribuído**: o voto se espalha de forma parecida pelas regiões — sem "
        "fortalezas marcantes, mas com presença em quase todo lugar. Capilaridade ampla.")
    add_p(doc,
        "**Híbrido**: tem fortalezas claras em algumas regiões, mas ainda chega em todas "
        "as outras com presença razoável. Reduto definido com cobertura de território.")
    add_p(doc,
        "**Concentrado**: o voto se prende a poucas regiões — em algumas é forte, em "
        "outras quase desaparece. Candidato de nicho territorial.")

    add_h2(doc, "2.4 As 5 zonas territoriais (Reposicionamento)")
    add_p(doc,
        "Quando um candidato muda de cargo, comparamos a votação dele no cargo origem com "
        "a votação de uma referência (alguém forte no cargo destino). Cada região é "
        "classificada de acordo com as duas Performances:")
    add_p(doc,
        "**Base compartilhada** — candidato e referência foram fortes ali (Performance "
        "positiva nos dois). Voto naturalmente compartilhado.")
    add_p(doc,
        "**Voto pessoal** — só o candidato foi forte. Voto puxado pelo nome dele, não "
        "pela base do cargo destino.")
    add_p(doc,
        "**Espaço a conquistar** — só a referência foi forte. Eleitorado da referência "
        "ainda não acessado pelo candidato.")
    add_p(doc,
        "**Terreno aberto** — nenhum dos dois foi forte, mas a região é grande em volume "
        "eleitoral. Eleitor numeroso, sem dono claro.")
    add_p(doc,
        "**Volume baixo** — nenhum dos dois foi forte e a região tem pouco eleitorado.")

    add_h2(doc, "2.5 Os 3 cenários")
    add_p(doc,
        "**Substituto orgânico** — cenário em que o candidato mantém apenas o eleitorado "
        "que já tinha (Base compartilhada + Voto pessoal), sem capturar nada da referência.")
    add_p(doc,
        "**Ponte parcial** — Substituto orgânico + 30% do eleitorado da referência no "
        "Espaço a conquistar.")
    add_p(doc,
        "**Ponte construída** — Substituto orgânico + 60% do eleitorado da referência no "
        "Espaço a conquistar. Cenário máximo do modelo.")
    add_p(doc,
        "Os percentuais (30% / 60%) são parâmetros canônicos. O cenário real depende de "
        "variáveis fora do escopo do modelo — campanha, alianças, conjuntura, performance "
        "dos adversários.")

    doc.add_page_break()

    # ── 3. PERFIL ELEITORAL DO CANDIDATO ──
    add_h1(doc, f"3. Perfil eleitoral do candidato em 2022")

    add_h2(doc, "3.1 Trajetória")
    eleito_txt = "foi eleito" if origem.get("eleito") else "foi candidato a"
    add_p(doc,
        f"{nome} {eleito_txt} {orig_cargo_lbl} em 2022 com **{fmt_int(origem['total'])} "
        f"votos**, pelo {origem.get('partido', '?')}, alinhado ao campo "
        f"{campo_lbl.lower()}. A votação se distribuiu entre as 28 Regiões Administrativas "
        f"do DF com voto registrado, com diferenças marcantes entre as regiões mais fortes "
        f"e as regiões mais fracas — o que define a partida da análise.")

    add_h2(doc, "3.2 Onde foi forte")
    nomes_forcas = ", ".join(r["ra"] for r in forcas)
    add_p(doc,
        f"As três regiões em que {nome_curto} teve a maior **Performance** foram "
        f"**{nomes_forcas}**. Em todas, recebeu bem mais voto do que a média do cargo "
        f"esperaria — são as fortalezas territoriais.")
    add_caption_tabela(doc, ctr, f"Maiores Performances de {nome_curto} em 2022")
    add_table(doc,
        ["Região", "Votos", "Performance", "Status"],
        [[r["ra"], fmt_int(r.get("v") or 0), fmt_idx(r.get("idx")),
          (r.get("s") or "—").replace("CAMPO MEDIO", "Esperado")
                              .replace("CAMPO MÉDIO", "Esperado").title()]
         for r in forcas])
    perfil_forcas_desc = PERFIL_DESC.get(perfil_forcas, "regiões mistas")
    soma_top = sum((r.get("v") or 0) for r in forcas)
    pct_top = soma_top / origem["total"] * 100 if origem["total"] else 0
    add_p(doc,
        f"Essas três regiões somam {fmt_int(soma_top)} votos — {fmt_pct(pct_top)} do total "
        f"de {nome_curto} em 2022. Predominam **{perfil_forcas_desc}**.")

    add_h2(doc, "3.3 Onde foi fraco")
    nomes_fraq = ", ".join(r["ra"] for r in fraquezas)
    add_p(doc,
        f"No outro extremo, em **{nomes_fraq}**, {nome_curto} recebeu Performance "
        f"fortemente negativa — quer dizer, votação bem abaixo da média do cargo nessas "
        f"regiões.")
    add_caption_tabela(doc, ctr, f"Menores Performances de {nome_curto} em 2022")
    add_table(doc,
        ["Região", "Votos", "Performance", "Status"],
        [[r["ra"], fmt_int(r.get("v") or 0), fmt_idx(r.get("idx")),
          (r.get("s") or "—").replace("CAMPO MEDIO", "Esperado")
                              .replace("CAMPO MÉDIO", "Esperado").title()]
         for r in fraquezas])
    perfil_fraq_desc = PERFIL_DESC.get(perfil_fraq, "regiões mistas")
    add_p(doc,
        f"Essas três regiões caracterizam-se principalmente como **{perfil_fraq_desc}**. "
        f"São territórios onde o nome dele praticamente não pegou em 2022.")

    add_h2(doc, "3.4 Perfil de votação do candidato")
    if tip_o:
        add_p(doc, descrever_perfil_votacao(tip_o, nome_curto))

    if imgs.get("mapa_origem"):
        doc.add_picture(str(imgs["mapa_origem"]), width=Inches(6.0))
        add_caption_figura(doc, ctr,
            f"Performance de {nome} por Região Administrativa "
            f"({orig_cargo_lbl} 2022)")

    doc.add_page_break()

    # ── 4. COMO O DF VOTA ──
    add_h1(doc, "4. Como o DF vota: implicações para o campo")

    add_h2(doc, "4.1 A geografia do voto")
    total_aptos = sum((dados_ra[n].get("el_aptos") or 0) for n in dados_ra)
    ras_ord = sorted(dados_ra.items(), key=lambda kv: -(kv[1].get("el_aptos") or 0))
    top3_aptos = ras_ord[:3]
    top3_concent = sum((d.get("el_aptos") or 0) for _, d in top3_aptos) / total_aptos * 100
    add_p(doc,
        f"O DF tem {len(dados_ra)} Regiões Administrativas e cerca de {fmt_int(total_aptos)} "
        f"eleitores aptos. A concentração eleitoral é fortemente desigual: três regiões — "
        f"**{', '.join(n for n, _ in top3_aptos)}** — somam {fmt_pct(top3_concent)} de todo "
        f"o eleitorado do DF.")
    add_p(doc,
        "Em cargos majoritários (Governador, Senador), ganhar nas três grandes regiões é "
        "praticamente obrigatório. Em cargos proporcionais (Deputado), o candidato pode "
        "escolher entre concentrar voto em poucas regiões ou pulverizar — ambas as "
        "estratégias têm precedentes na história eleitoral do DF.")

    add_h2(doc, "4.2 O paradoxo do voto progressista")
    add_p(doc,
        "No DF acontece um padrão diferente do resto do Brasil. As regiões de maior renda "
        "— Plano Piloto, Lago Sul, Sudoeste, Jardim Botânico, Águas Claras — costumam "
        "votar mais em candidatos progressistas. No padrão nacional, esse mesmo perfil "
        "socioeconômico tende a votar no campo conservador.")
    add_p(doc,
        "A explicação tem dois lados. Um é o servidor federal: alta concentração no DF, "
        "com identificação de carreira pública e tendência ao voto progressista. O outro "
        "é a classe média alta privada do eixo Lago Sul / Sudoeste / Jardim Botânico / "
        "Águas Claras, com escolaridade alta e perfil de voto também progressista. Os "
        "dois grupos somados desenham um teto urbano de classe média/alta majoritariamente "
        "progressista no DF — algo único no panorama nacional.")

    add_h2(doc, "4.3 Implicação para o campo do candidato")
    add_p(doc, implicacao_paradoxo_campo(campo, cargo_dest))

    if imgs.get("scatter"):
        doc.add_picture(str(imgs["scatter"]), width=Inches(6.0))
        add_caption_figura(doc, ctr,
            "Relação entre presença de servidor federal na população (PDAD) "
            "e voto progressista para Deputado Federal (TSE 2022) — uma RA por ponto")

    doc.add_page_break()

    # ── 5. A REFERÊNCIA DESTINO ──
    add_h1(doc, "5. A referência destino")

    add_h2(doc, "5.1 Trajetória eleitoral da referência")
    eleito_ref_txt = "foi eleito" if ref.get("eleito") else "foi candidato a"
    add_p(doc,
        f"{ref_nome} {eleito_ref_txt} {cargo_dest_lbl} em 2022 com "
        f"**{fmt_int(ref['total'])} votos** pelo {ref.get('partido', '?')}, alinhado ao "
        f"campo {CAMPO_LBL.get(ref['campo'], ref['campo']).lower()}. Esse total ficou "
        f"{ref_vs_meta}.")

    add_h2(doc, "5.2 A escala da referência")
    add_p(doc,
        f"O eleitorado da referência tem escala muito maior do que o do candidato: "
        f"**{fmt_int(ref['total'])} votos** contra **{fmt_int(origem['total'])}** de "
        f"{nome_curto} em 2022 — uma razão de **{razao_ref_str}x**. Em outras palavras, "
        f"a referência tem cerca de {razao_ref_str} vezes o tamanho do eleitorado do "
        f"candidato.")
    add_p(doc,
        f"Em {n_ras_ref_pos} das 28 RAs com voto registrado, {ref_nome_curto} teve "
        f"Performance positiva (recebeu mais voto que a média do cargo). É um indicador "
        f"de quão amplo é o território onde a referência foi forte.")

    add_h2(doc, "5.3 Onde a referência foi forte (por Performance)")
    add_p(doc,
        f"As três regiões em que {ref_nome_curto} teve maior Performance foram "
        f"**{', '.join(r['ra'] for r in ref_forcas_idx)}**.")
    add_caption_tabela(doc, ctr, f"Maiores Performances de {ref_nome_curto} em 2022")
    add_table(doc,
        ["Região", "Votos", "Performance"],
        [[r["ra"], fmt_int(r.get("v") or 0), fmt_idx(r.get("idx"))]
         for r in ref_forcas_idx])
    perfil_ref_desc = PERFIL_DESC.get(perfil_ref_forcas, "regiões mistas")
    add_p(doc,
        f"Essas regiões predominam em **{perfil_ref_desc}** — é onde o eleitor da "
        f"referência se concentra com mais intensidade.")

    add_h2(doc, "5.4 Onde a referência foi forte em volume absoluto")
    add_p(doc,
        f"Em alguns casos, a região onde a referência teve mais Performance não é a "
        f"mesma onde teve mais votos absolutos. Em RAs grandes, mesmo Performance "
        f"moderada se traduz em volume substantivo.")
    add_caption_tabela(doc, ctr, f"Maiores volumes absolutos de votos de {ref_nome_curto} em 2022")
    add_table(doc,
        ["Região", "Votos", "Performance", "Aptos da RA"],
        [[r["ra"], fmt_int(r.get("v") or 0), fmt_idx(r.get("idx")),
          fmt_int((dados_ra.get(r["ra"]) or {}).get("el_aptos") or 0)]
         for r in ref_forcas_vol])

    add_h2(doc, "5.5 Onde a referência foi fraca")
    add_p(doc,
        f"Nas três regiões com pior Performance, "
        f"**{', '.join(r['ra'] for r in ref_fraquezas)}**, a referência ficou abaixo da "
        f"média do cargo. São territórios onde, mesmo sendo a candidatura dominante no "
        f"DF em volume, a referência teve menor penetração relativa.")
    add_caption_tabela(doc, ctr, f"Menores Performances de {ref_nome_curto} em 2022")
    add_table(doc,
        ["Região", "Votos", "Performance"],
        [[r["ra"], fmt_int(r.get("v") or 0), fmt_idx(r.get("idx"))]
         for r in ref_fraquezas])

    add_h2(doc, "5.6 Perfil de votação da referência")
    if tip_r:
        add_p(doc, descrever_perfil_votacao(tip_r, ref_nome_curto))
    if tip_o and tip_r:
        if tip_o["lbl"] == tip_r["lbl"]:
            comparativo = (f"Os Perfis de votação de {nome_curto} e {ref_nome_curto} são "
                           f"semelhantes — ambos {tip_o['lbl']}. Capilaridade comparável.")
        elif tip_o["lbl"] == "Concentrado" and tip_r["lbl"] == "Distribuído":
            comparativo = (f"{nome_curto} é Concentrado e {ref_nome_curto} é Distribuído — "
                           f"perfis bem diferentes. A campanha de captura precisa lidar com "
                           f"essa diferença: o candidato puxa nicho, a referência puxa rede.")
        elif tip_o["lbl"] == "Distribuído" and tip_r["lbl"] == "Concentrado":
            comparativo = (f"{nome_curto} é Distribuído e {ref_nome_curto} é Concentrado — "
                           f"o candidato cobre mais território, mas a referência tem peso "
                           f"em poucas regiões específicas.")
        else:
            comparativo = (f"{nome_curto} é {tip_o['lbl']} e {ref_nome_curto} é "
                           f"{tip_r['lbl']} — capilaridades distintas.")
        add_p(doc, comparativo)

    add_h2(doc, "5.7 O eleitorado da referência por perfil de região")
    perfil_voto_ref = {}
    for r in ras_ref:
        if r.get("idx") is None: continue
        p = PERFIL_RA.get(r["ra"], "outro")
        if p == "outro": continue
        perfil_voto_ref[p] = perfil_voto_ref.get(p, 0) + (r.get("v") or 0)
    total_ref = sum(perfil_voto_ref.values()) or 1
    perfil_table_rows = []
    for p_key in ["AB", "media", "popular", "rural"]:
        v = perfil_voto_ref.get(p_key, 0)
        perfil_table_rows.append([
            PERFIL_DESC.get(p_key, p_key).capitalize(),
            fmt_int(v),
            fmt_pct(v / total_ref * 100),
        ])
    add_caption_tabela(doc, ctr,
        f"Distribuição dos votos de {ref_nome_curto} por perfil de região")
    add_table(doc,
        ["Perfil de região", "Votos da referência", "% do total da referência"],
        perfil_table_rows)
    perfil_dom_ref = max(perfil_voto_ref, key=perfil_voto_ref.get) if perfil_voto_ref else None
    if perfil_dom_ref:
        pct_dom_ref = perfil_voto_ref[perfil_dom_ref] / total_ref * 100
        add_p(doc,
            f"O eleitorado de {ref_nome_curto} se concentra principalmente em "
            f"**{PERFIL_DESC.get(perfil_dom_ref, perfil_dom_ref)}** "
            f"({fmt_pct(pct_dom_ref)} do total). É um retrato do tipo de eleitor que a "
            f"campanha de captura precisaria conversar para herdar parte dessa base.")

    add_h2(doc, "5.8 Mapa da referência por Performance")
    if imgs.get("mapa_referencia"):
        doc.add_picture(str(imgs["mapa_referencia"]), width=Inches(6.0))
        add_caption_figura(doc, ctr,
            f"Performance de {ref_nome} por Região Administrativa "
            f"({cargo_dest_lbl} 2022)")

    add_h2(doc, "5.9 O que isso significa para o reposicionamento")
    add_p(doc, perfil_intersecao(cats))
    add_p(doc,
        f"Considerando os dois eleitorados juntos: o de {nome_curto} ({fmt_int(origem['total'])} "
        f"votos) e o de {ref_nome_curto} ({fmt_int(ref['total'])} votos), a campanha de "
        f"reposicionamento opera dentro do território onde os dois se cruzam ou onde apenas "
        f"a referência foi forte. Capturar uma fração desse eleitorado da referência — "
        f"mesmo que pequena — pode ter peso significativo na soma final, dada a diferença "
        f"de escala entre os dois.")

    doc.add_page_break()

    # ── 6. A GEOMETRIA DA CAPTURA: AS 5 ZONAS ──
    add_h1(doc, "6. A geometria da captura: as 5 zonas territoriais")
    add_p(doc,
        f"Cruzando a Performance de {nome_curto} no cargo origem com a Performance de "
        f"{ref_nome_curto} no cargo destino, cada uma das 28 regiões cai em uma das cinco "
        f"zonas. O agrupamento desenha o terreno onde a campanha vai operar.")

    if imgs.get("mapa_categorias"):
        doc.add_picture(str(imgs["mapa_categorias"]), width=Inches(6.5))
        add_caption_figura(doc, ctr,
            "Distribuição das Regiões Administrativas do DF nas 5 zonas territoriais "
            "do reposicionamento")

    add_caption_tabela(doc, ctr, "Distribuição das RAs e do eleitorado nas 5 zonas")
    add_table(doc,
        ["Zona territorial", "Nº de RAs", "Aptos somados"],
        [[CAT_DEF[k]["lbl"], str(len(cats[k])),
          fmt_int(sum(r["aptos"] for r in cats[k]))] for k in CAT_ORDER])

    # 6.1 Base compartilhada
    add_h2(doc, "6.1 Base compartilhada")
    if arr_comp:
        nomes_comp = ", ".join(r["ra"] for r in arr_comp[:6]) + (f" e mais {len(arr_comp)-6}" if len(arr_comp) > 6 else "")
        add_p(doc,
            f"São **{len(arr_comp)} regiões** onde {nome_curto} foi forte e "
            f"{ref_nome_curto} também: {nomes_comp}. O voto aqui é naturalmente "
            f"compartilhado pelos dois — o eleitor é compatível com os dois perfis. "
            f"O candidato já tem **{fmt_int(votos_comp)} votos próprios** "
            f"({fmt_pct(votos_comp/meta*100 if meta else 0)} do Patamar de eleição) "
            f"garantidos como piso, desde que a operação se mantenha.")
        add_caption_tabela(doc, ctr, "Regiões em Base compartilhada")
        add_table(doc,
            ["Região", "Aptos", "Votos do candidato", "Votos da referência"],
            [[r["ra"], fmt_int(r["aptos"]), fmt_int(r["orig_votos"]), fmt_int(r["ref_votos"])]
             for r in arr_comp])
    else:
        add_p(doc, f"Nenhuma região em Base compartilhada — {nome_curto} e {ref_nome_curto} "
                   f"não foram fortes simultaneamente em nenhuma RA.")

    # 6.2 Voto pessoal
    add_h2(doc, "6.2 Voto pessoal")
    if arr_nuc:
        add_p(doc,
            f"São **{len(arr_nuc)} regiões** onde {nome_curto} foi forte mas "
            f"{ref_nome_curto} foi fraca: **{', '.join(r['ra'] for r in arr_nuc)}**. Aqui "
            f"o voto vem do nome do candidato — não do eleitorado típico do cargo destino. "
            f"Total: **{fmt_int(votos_nuc)} votos**.")
        add_caption_tabela(doc, ctr, "Regiões em Voto pessoal")
        add_table(doc,
            ["Região", "Aptos", "Votos do candidato", "Votos da referência"],
            [[r["ra"], fmt_int(r["aptos"]), fmt_int(r["orig_votos"]), fmt_int(r["ref_votos"])]
             for r in arr_nuc])
    else:
        add_p(doc,
            f"Nenhuma região isolada em Voto pessoal. Onde {nome_curto} foi forte, "
            f"{ref_nome_curto} também foi — toda a base do candidato é compartilhada com a "
            f"referência. Não há \"reduto exclusivo\" do candidato fora da Base compartilhada.")

    # 6.3 Espaço a conquistar
    add_h2(doc, "6.3 Espaço a conquistar")
    if arr_lac:
        add_p(doc,
            f"São **{len(arr_lac)} regiões** onde {ref_nome_curto} foi forte mas "
            f"{nome_curto} não chegou: **{', '.join(r['ra'] for r in arr_lac)}**. "
            f"Predominam **{perfil_lac_desc}**. **Esses {fmt_int(lacuna_total)} votos** "
            f"compõem o universo total que o modelo considera capturável a partir da "
            f"referência.")
        add_caption_tabela(doc, ctr, "Regiões em Espaço a conquistar")
        add_table(doc,
            ["Região", "Aptos", "Votos do candidato", "Votos da referência"],
            [[r["ra"], fmt_int(r["aptos"]), fmt_int(r["orig_votos"]), fmt_int(r["ref_votos"])]
             for r in arr_lac])
    else:
        add_p(doc,
            f"Nenhuma região em Espaço a conquistar. {ref_nome_curto} não foi forte em "
            f"regiões onde {nome_curto} foi fraco — não há universo de captura disponível "
            f"pelo modelo.")

    # 6.4 Terreno aberto
    add_h2(doc, "6.4 Terreno aberto")
    if arr_hos:
        add_p(doc,
            f"São **{len(arr_hos)} regiões** com volume eleitoral alto "
            f"({fmt_int(aptos_hos)} aptos somados) onde nem {nome_curto} nem "
            f"{ref_nome_curto} foram fortes em 2022: "
            f"**{', '.join(r['ra'] for r in arr_hos)}**. É território com muito eleitor "
            f"em aberto — sem dono.")
    else:
        add_p(doc, "Nenhuma região classificada como Terreno aberto.")

    # 6.5 Volume baixo
    add_h2(doc, "6.5 Volume baixo")
    if arr_per:
        add_p(doc,
            f"São **{len(arr_per)} regiões** sem base e com pouco eleitorado "
            f"({fmt_int(aptos_per)} aptos somados): "
            f"**{', '.join(r['ra'] for r in arr_per)}**.")
    else:
        add_p(doc, "Nenhuma região classificada como Volume baixo.")

    doc.add_page_break()

    # ── 7. A META E OS TRÊS CENÁRIOS ──
    add_h1(doc, "7. A meta e os três cenários")

    add_h2(doc, "7.1 O Patamar de eleição")
    eh_maj = cargo_dest in ("GOVERNADOR", "SENADOR")
    if eh_maj:
        add_p(doc,
            f"O **Patamar de eleição** para {cargo_dest_lbl} é **{fmt_int(meta)} votos** "
            f"— volume mínimo histórico para chegar ao 2º turno (Governador) ou ser eleito "
            f"(Senador), com base nos resultados de 2022.")
    else:
        add_p(doc,
            f"O **Patamar de eleição** para {cargo_dest_lbl} é **{fmt_int(meta)} votos** "
            f"— faixa onde candidatos individuais começam a ser efetivamente eleitos, com "
            f"base na distribuição dos eleitos em 2022.")
    add_p(doc,
        f"{nome_curto} parte de **{fmt_int(origem['total'])} votos** em 2022 — "
        f"{fmt_pct(pct_origem_meta)} desse Patamar. O salto é de **{salto_str}x** o "
        f"eleitorado atual.")

    add_h2(doc, "7.2 Cenário Substituto orgânico")
    add_p(doc,
        f"Cenário base: o candidato mantém apenas o eleitorado que já tinha — Base "
        f"compartilhada + Voto pessoal — sem capturar nada da referência. Resultado: "
        f"**{fmt_int(substituto)} votos** ({fmt_pct(pct_sub)} do Patamar).")

    add_h2(doc, "7.3 Cenário Ponte parcial")
    add_p(doc,
        f"Substituto + 30% dos votos de {ref_nome_curto} no Espaço a conquistar "
        f"({fmt_int(round(0.30 * lacuna_total))} votos capturados). Resultado: "
        f"**{fmt_int(pcl)} votos** ({fmt_pct(pct_pcl)} do Patamar).")

    add_h2(doc, "7.4 Cenário Ponte construída")
    add_p(doc,
        f"Cenário máximo do modelo: Substituto + 60% do Espaço a conquistar "
        f"({fmt_int(round(0.60 * lacuna_total))} votos capturados). Resultado: "
        f"**{fmt_int(cns)} votos** ({fmt_pct(pct_cns)} do Patamar).")

    add_caption_tabela(doc, ctr, "Os três cenários do modelo de Reposicionamento")
    add_table(doc,
        ["Cenário", "Votos projetados", "% do Patamar", "Captura"],
        [
            ["Substituto orgânico", fmt_int(substituto), fmt_pct(pct_sub), "0% do Espaço a conquistar"],
            ["Ponte parcial",       fmt_int(pcl),        fmt_pct(pct_pcl), "30% do Espaço a conquistar"],
            ["Ponte construída",    fmt_int(cns),        fmt_pct(pct_cns), "60% do Espaço a conquistar"],
        ])

    add_h2(doc, "7.5 O que os números mostram")
    add_block_note(doc, "", leitura)

    doc.add_page_break()

    # ── 8. IMPLICAÇÕES ESTRATÉGICAS ──
    add_h1(doc, "8. Implicações estratégicas")
    add_p(doc,
        "Cada uma das 5 zonas territoriais sugere um tipo distinto de operação. As "
        "decisões finais — quanto investir, quando, com qual mensagem — dependem de "
        "variáveis fora do escopo desta análise (campanha, alianças, conjuntura, "
        "performance dos adversários).")

    add_h2(doc, "8.1 Defender a Base compartilhada")
    add_p(doc,
        f"As **{len(arr_comp)} regiões** de Base compartilhada garantem "
        f"**{fmt_int(votos_comp)} votos** ({fmt_pct(votos_comp/meta*100 if meta else 0)} do "
        f"Patamar) — desde que a operação se mantenha. Voto compartilhado entre "
        f"{nome_curto} e {ref_nome_curto}; toda perda aqui é perda líquida na soma.")

    add_h2(doc, "8.2 Capturar o Espaço a conquistar")
    if arr_lac:
        add_p(doc,
            f"As **{len(arr_lac)} regiões** do Espaço a conquistar "
            f"({', '.join(r['ra'] for r in arr_lac[:5])}) oferecem "
            f"**{fmt_int(lacuna_total)} votos** disponíveis. A mensagem da campanha "
            f"precisa conectar especificamente com o eleitorado desse perfil — "
            f"**{perfil_lac_desc}** — para que a captura aconteça.")
    else:
        add_p(doc, "Sem Espaço a conquistar disponível pelo modelo nesta configuração.")

    add_h2(doc, "8.3 Decidir sobre Terreno aberto")
    if arr_hos:
        add_p(doc,
            f"As **{len(arr_hos)} regiões** de Terreno aberto concentram "
            f"**{fmt_int(aptos_hos)} aptos** sem base do candidato nem da referência. "
            f"É a maior reserva de eleitores fora do alcance direto do modelo — entrar "
            f"aqui depende de variáveis externas (mensagem, conjuntura, alianças). "
            f"A campanha precisa decidir se tem condições de operar nessas regiões "
            f"sem comprometer o piso da Base compartilhada e a captura do Espaço a conquistar.")
    else:
        add_p(doc, "Sem regiões em Terreno aberto.")

    add_h2(doc, "8.4 Fidelizar o Voto pessoal")
    if arr_nuc:
        add_p(doc,
            f"As **{len(arr_nuc)} regiões** de Voto pessoal "
            f"({', '.join(r['ra'] for r in arr_nuc)}) guardam **{fmt_int(votos_nuc)} "
            f"votos** próprios — voto puxado pelo nome do candidato, não pela base do "
            f"cargo destino. Sem trabalho de campo específico, esses votos podem migrar "
            f"para outros candidatos do mesmo perfil pessoal.")
    else:
        add_p(doc,
            f"Sem regiões em Voto pessoal isolado — toda a base de {nome_curto} já é "
            f"compartilhada com {ref_nome_curto}.")

    add_h2(doc, "8.5 Volume baixo: presença mínima")
    add_p(doc,
        f"As **{len(arr_per)} regiões** de Volume baixo ({fmt_int(aptos_per)} aptos) "
        f"não têm peso relevante na soma final. Recebem presença mínima — qualquer real "
        f"investido aqui rende menos do que em qualquer outra zona.")

    add_h2(doc, "8.6 O que os números mostram")
    add_block_note(doc, "Síntese dos cenários:", leitura)
    if pct_cns < 60:
        add_p(doc,
            "Possibilidades sobre a mesa quando o Espaço a conquistar é menor do que "
            "a diferença até o Patamar de eleição:")
        if arr_hos:
            add_p(doc,
                f"**a)** Operar no Terreno aberto — {fmt_int(aptos_hos)} aptos em "
                f"{len(arr_hos)} regiões "
                f"({', '.join(r['ra'] for r in arr_hos[:4])}). Custos e risco proporcionais.")
        else:
            add_p(doc,
                "**a)** Expansão para territórios fora do alcance do modelo (eleitorado novo, "
                "mensagem nova).")
        add_p(doc,
            f"**b)** Coligação com candidato de capilaridade complementar — outro nome do "
            f"mesmo campo cuja base territorial preencha as regiões onde {nome_curto} foi "
            f"fraco. Captura por aliança em vez de por mensagem.")
        add_p(doc,
            "**c)** Reavaliação do cargo destino — outros cargos têm Patamar de eleição e "
            "geometria distintas; vale simular o diagnóstico no cargo alternativo antes de "
            "comprometer-se com este.")
        add_p(doc,
            "Cada uma dessas possibilidades tem custos e riscos próprios, fora do escopo "
            "desta análise. Como pesar essas opções é decisão da campanha.")

    doc.save(output_path)


# ── API pública ───────────────────────────────────────────────────────────────

def gerar(config, dados_preloaded, output, log=print):
    """
    Gera o relatório Word.

    Args:
      config: dict no formato esperado (tipo, origem, destino, referencia)
      dados_preloaded: {"cands": [...], "dados_ra": {...}} já carregados
                       (evita recarregar o CSV grande do TSE em cada chamada)
      output: pathlib.Path ou file-like (BytesIO) — destino do .docx
      log: função de logging (default: print)

    Returns: dict com metadados {origem, ref, cats_counts, slug}.
    """
    if config.get("tipo") != "reposicionamento":
        raise ValueError(f"Tipo não suportado: {config.get('tipo')}")

    cands = dados_preloaded["cands"]
    dados_ra = dados_preloaded["dados_ra"]

    origem = buscar_candidato(cands, config["origem"]["nome"], config["origem"]["cargo"])
    ref = buscar_candidato(cands, config["referencia"]["nome"], config["referencia"]["cargo"])
    log(f"  origem: {origem['nome']} ({CARGO_LBL[origem['cargo']]})")
    log(f"  ref:    {ref['nome']} ({CARGO_LBL[ref['cargo']]})")

    cats, med_aptos, total_ras = computar_categorias(origem, ref, dados_ra)

    tmpdir = Path(tempfile.mkdtemp(prefix="diag_"))
    imgs = {}
    try:
        scatter_p = tmpdir / "scatter.png"
        if gerar_scatter_png(dados_ra, scatter_p): imgs["scatter"] = scatter_p
        gdf = carregar_geo()
        mapa_o = tmpdir / "mapa_origem.png"
        gerar_mapa_idx(origem, gdf, mapa_o, "")  # título via caption
        imgs["mapa_origem"] = mapa_o
        mapa_r = tmpdir / "mapa_referencia.png"
        gerar_mapa_idx(ref, gdf, mapa_r, "")
        imgs["mapa_referencia"] = mapa_r
        mapa_c = tmpdir / "mapa_categorias.png"
        gerar_mapa_categorias(cats, gdf, mapa_c)
        imgs["mapa_categorias"] = mapa_c
        montar_word(config, origem, ref, dados_ra, cats, med_aptos, total_ras, imgs, output)
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

    return {
        "origem_nome": origem["nome"],
        "ref_nome": ref["nome"],
        "cats_counts": {k: len(cats[k]) for k in CAT_ORDER},
        "slug": slugify(origem["nome"]),
    }


def carregar_dados():
    """Helper: carrega tudo que `gerar()` precisa pré-carregado."""
    df_ipe, df_mestre, df_narr, df_campo, df_cand = carregar()
    dados_ra = montar_dados(df_ipe, df_mestre, df_narr, df_campo, df_cand)
    cands = montar_candidatos()
    return {"dados_ra": dados_ra, "cands": cands}


# ── Main (CLI) ────────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 2:
        print("Uso: python3 gerar_relatorio_diag.py diag_config.json")
        sys.exit(1)
    cfg_path = Path(sys.argv[1])
    if not cfg_path.exists():
        print(f"Arquivo não encontrado: {cfg_path}"); sys.exit(1)
    config = json.loads(cfg_path.read_text(encoding="utf-8"))

    print("\n  gerar_relatorio_diag.py")
    print("  " + "─" * 40)
    print("\n  [1/3] Carregando dados...")
    dados = carregar_dados()
    print("  [2/3] Gerando relatório...")
    slug = slugify(config["origem"]["nome"])
    out = Path(f"relatorio_diag_{slug}.docx")
    meta = gerar(config, dados, out, log=lambda m: print("       " + m))
    print(f"  [3/3] Salvo")
    print(f"\n  ✓ {out}  ({out.stat().st_size // 1024} KB)")
    for k in CAT_ORDER:
        print(f"     {CAT_DEF[k]['lbl']:<25} {meta['cats_counts'][k]:>2} RAs")
    print()


if __name__ == "__main__":
    main()
